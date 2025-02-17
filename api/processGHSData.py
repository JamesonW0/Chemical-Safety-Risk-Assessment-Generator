import io
from datetime import datetime
from copy import deepcopy
import docx
from flask import Flask, request, send_file
import os
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ----- Efficiency improvements: preload template files into memory -----
FORM_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "COSHH_Form_Template.docx")
TICKS_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "COSHH_Ticks_Template.docx")

def load_form_template():
    return docx.Document(FORM_TEMPLATE_PATH)

def load_ticks_template():
    return docx.Document(TICKS_TEMPLATE_PATH)

# ----- Global mappings for hazard codes, exposure routes, and control measures -----
hazard_mappings = {
    "spill": {200, 201, 202, 203, 204, 205, 206, 207, 208, 230, 231, 232, 250, 251, 300, 304, 310, 330, 340, 301, 311, 331},
    "flame": {200, 201, 202, 203, 204, 205, 206, 207, 208, 220, 221, 222, 223, 224, 225, 226, 227, 228, 229, 230, 231, 232,
              240, 241, 242, 251, 252, 270, 271, 272},
    "Tcontrol": {200, 201, 202, 203, 204, 205, 206, 207, 208, 225, 226, 227, 228, 230, 231, 270, 271, 272, 280},
    "pregnant": {360, 361, 362},
    "water": {261, 262},
    "dropwise": {261, 262, 270, 271, 272},
    "air": {230, 231, 232, 250},
}

exposure_mappings = {
    "eye": {314, 318, 319},
    "skin": {310, 311, 312, 314, 315, 317},
    "inhalation": {304, 330, 331, 332, 334, 335, 336},
    "ingestion": {300, 301, 302, 304},
}

def get_routes_and_measures(hazard_statements):
    """
    Process a newline-separated string of hazard codes (as text)
    and return exposure routes (list of 4 ints) and control measures (list of 9 ints).
    """
    if not hazard_statements or hazard_statements[0] == "N":
        return [0] * 4, [0] * 9

    hazard_codes = set()
    for row in hazard_statements.splitlines():
        try:
            code = int(row.strip())
        except Exception:
            try:
                code = int(row[1:4])
            except Exception:
                continue
        hazard_codes.add(code)

    exp_routes = [0] * 4
    ctrl_measures = [1 if i == 2 else 0 for i in range(9)]

    for key, hazard_set in hazard_mappings.items():
        if hazard_codes & hazard_set:
            index = {"spill": 0, "flame": 3, "Tcontrol": 4, "pregnant": 5,
                     "water": 6, "dropwise": 7, "air": 8}.get(key)
            if index is not None:
                ctrl_measures[index] = 1

    for key, hazard_set in exposure_mappings.items():
        if hazard_codes & hazard_set:
            index = {"eye": 0, "skin": 1, "inhalation": 2, "ingestion": 3}.get(key)
            if index is not None:
                exp_routes[index] = 1

    return exp_routes, ctrl_measures

def prepare_new_form():
    """
    Open the COSHH form template and update the date.
    """
    doc_ = load_form_template()
    doc_.tables[0].rows[1].cells[3].text = datetime.today().strftime("%d/%m/%Y")
    return doc_

def change_COSHH_row_one(doc_, chem_name_, amount_, hazard_statements_):
    table = doc_.tables[2]
    table.rows[1].cells[0].text = chem_name_
    table.rows[1].cells[1].text = amount_
    table.rows[1].cells[2].text = hazard_statements_
    return doc_

def add_COSHH_row(doc_, chem_name_, amount_, hazard_statements_, exp_routes_, ctrl_measures_):
    tick_doc = load_ticks_template()
    table = doc_.tables[2]
    new_row = deepcopy(table.rows[1]._tr)
    table._tbl.append(new_row)
    new_row = table.rows[-1]
    new_row.cells[0].text = chem_name_
    new_row.cells[1].text = amount_
    new_row.cells[2].text = hazard_statements_
    # Set exposure routes
    temp_tbl = tick_doc.tables[0]
    new_row.cells[3]._element.clear()
    for i in range(4):
        new_row.cells[3]._element.extend(deepcopy(temp_tbl.rows[i].cells[exp_routes_[i]]._element))
    # Set control measures
    temp_tbl = tick_doc.tables[1]
    new_row.cells[4]._element.clear()
    for i in range(9):
        new_row.cells[4]._element.extend(deepcopy(temp_tbl.rows[i].cells[ctrl_measures_[i]]._element))
    return doc_

# ----- Flask Application Setup -----
app = Flask(__name__)

@app.route("/api/processGHSData", methods=["POST"])
def submit_ghs_data():
    try:
        # Receive JSON from the frontend.
        # Expected structure: [{"name": "Chem A", "amount": "x L", "hazards": ["200", "201"]}, {...}, ...]
        data = request.get_json()
        if not data:
            return {"error": "No data received"}, 400

        # Prepare a new COSHH form.
        doc = prepare_new_form()

        # Process additional chemicals by adding rows.
        for chem in data:
            hazard_str = "\n".join(chem.get("hazards", []))
            exp_routes, ctrl_measures = get_routes_and_measures(hazard_str)
            doc = add_COSHH_row(doc, chem.get("name", ""), chem.get("amount", ""), hazard_str, exp_routes, ctrl_measures)

        # Save the document to an in-memory stream.
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = f"COSHH.docx"
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.exception("Error processing COSHH data")
        return {"error": str(e)}, 500

# For local testing only; in production, Vercel will handle the function invocation.
if __name__ == "__main__":
    app.run()
